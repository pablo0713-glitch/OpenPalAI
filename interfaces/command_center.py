from __future__ import annotations

import base64
import io
import json
import mimetypes
import re
import time
from collections import deque
from pathlib import Path
from typing import Any

from fastapi import APIRouter, File, Form, UploadFile
from fastapi.responses import HTMLResponse, JSONResponse, RedirectResponse
from pydantic import BaseModel

from config.settings import Settings
from core.agent import AgentCore
from core.persona import (
    MessageContext,
    _normalize_agent_id,
    get_agent_config,
    get_companion_agent,
    get_default_agent_id,
    list_companion_agents,
    resolve_platform_agent_id,
)
from memory.library_store import LibraryStore

_ROOT = Path(__file__).parent.parent
_COMMAND_DIR = _ROOT / "command"
_MAX_UPLOAD_BYTES = 5 * 1024 * 1024
_MAX_TEXT_DOC_CHARS = 16000
_MAX_LIBRARY_DOC_CHARS = 60000
_MAX_EMBEDDED_IMAGES = 6
_TEXT_MEDIA_TYPES = {
    "application/json",
    "application/ld+json",
    "application/xml",
    "application/x-yaml",
    "application/yaml",
}
_TEXT_EXTENSIONS = {
    ".md",
    ".markdown",
    ".txt",
    ".json",
    ".yaml",
    ".yml",
    ".xml",
    ".csv",
    ".py",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".html",
    ".css",
    ".sql",
    ".sh",
    ".log",
}
_SUPPORTED_IMAGE_MEDIA_TYPES = {
    "image/jpeg",
    "image/png",
    "image/gif",
    "image/webp",
}
_SUPPORTED_DOCUMENT_MEDIA_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def _asset_version(path: Path) -> int:
    try:
        return int(path.stat().st_mtime)
    except OSError:
        return 0


def create_command_center_router(agent: AgentCore, settings: Settings) -> APIRouter:
    router = APIRouter()

    class _LibraryUpdateBody(BaseModel):
        always_on: bool

    @router.get("/", include_in_schema=False)
    async def root_redirect() -> RedirectResponse:
        return RedirectResponse(url="/command", status_code=307)

    @router.get("/command")
    async def command_index() -> HTMLResponse:
        html = (_COMMAND_DIR / "index.html").read_text(encoding="utf-8")
        css_version = _asset_version(_COMMAND_DIR / "style.css")
        js_version = _asset_version(_COMMAND_DIR / "app.js")
        html = html.replace("/command/static/style.css", f"/command/static/style.css?v={css_version}")
        html = html.replace("/command/static/app.js", f"/command/static/app.js?v={js_version}")
        return HTMLResponse(html, headers={"Cache-Control": "no-store"})

    @router.get("/command/status")
    async def command_status(agent_id: str = "") -> JSONResponse:
        cfg = get_agent_config()
        active_agent_id = resolve_platform_agent_id("command", agent_id, require_selectable=True, cfg=cfg)
        agent_cfg = get_companion_agent(active_agent_id, cfg)
        model_name = settings.claude_model if settings.model_provider == "anthropic" else (
            settings.ollama_model if settings.model_provider == "ollama" else settings.openai_model
        )
        return JSONResponse(
            {
                "agent_id": active_agent_id,
                "default_agent_id": get_default_agent_id(cfg),
                "agent_name": agent_cfg.get("agent_name", "Agent"),
                "command_center_name": cfg.get("command_center_name", "Command Center"),
                "agent_profile_image": agent_cfg.get("agent_profile_image", ""),
                "agents": list_companion_agents("command", selectable_only=True, cfg=cfg),
                "model_provider": settings.model_provider,
                "model_name": model_name,
                "uploads": {
                    "images": True,
                    "documents": True,
                    "pdf": True,
                    "docx": True,
                    "max_upload_bytes": _MAX_UPLOAD_BYTES,
                },
            }
        )

    @router.post("/command/chat")
    async def command_chat(
        message: str = Form(""),
        conversation_id: str = Form("default"),
        command_user_id: str = Form(""),
        display_name: str = Form(""),
        agent_id: str = Form(""),
        files: list[UploadFile] | None = File(default=None),
    ) -> JSONResponse:
        cleaned_message = message.strip()
        cleaned_conversation_id = _command_channel_id(conversation_id)
        cleaned_command_user_id = _command_user_id(command_user_id)
        cfg = get_agent_config()
        selected_agent_id = resolve_platform_agent_id("command", agent_id, require_selectable=True, cfg=cfg)
        agent_cfg = get_companion_agent(selected_agent_id, cfg)
        cleaned_display_name = display_name.strip() or cfg.get("command_center_name", "Command Center")
        uploads = files or []

        if not cleaned_message and not uploads:
            return JSONResponse({"ok": False, "error": "Message or attachment required."}, status_code=400)

        try:
            content, attachments = await _build_user_content(cleaned_message, uploads)
        except ValueError as exc:
            return JSONResponse({"ok": False, "error": str(exc)}, status_code=400)

        context = MessageContext(
            platform="command",
            user_id=cleaned_command_user_id,
            channel_id=cleaned_conversation_id,
            display_name=cleaned_display_name,
            agent_id=selected_agent_id,
        )
        response = await agent.handle_message(content, context)

        return JSONResponse(
            {
                "ok": True,
                "agent_id": selected_agent_id,
                "agent_name": agent_cfg.get("agent_name", "Agent"),
                "agent_profile_image": agent_cfg.get("agent_profile_image", ""),
                "conversation_id": cleaned_conversation_id,
                "reply": response.text,
                "attachments": attachments,
                "rate_limited": response.was_rate_limited,
                "refused": response.was_refused,
            }
        )

    @router.get("/command/history")
    async def command_history(
        conversation_id: str,
        command_user_id: str,
        agent_id: str = "",
    ) -> JSONResponse:
        cleaned_conversation_id = _command_channel_id(conversation_id)
        cleaned_command_user_id = _command_user_id(command_user_id)
        cfg = get_agent_config()
        selected_agent_id = resolve_platform_agent_id("command", agent_id, require_selectable=True, cfg=cfg)
        agent_cfg = get_companion_agent(selected_agent_id, cfg)
        history = await agent.get_conversation_history(
            cleaned_command_user_id,
            cleaned_conversation_id,
            agent_id=selected_agent_id,
        )
        return JSONResponse(
            {
                "ok": True,
                "agent_id": selected_agent_id,
                "agent_name": agent_cfg.get("agent_name", "Agent"),
                "conversation_id": cleaned_conversation_id,
                "command_user_id": cleaned_command_user_id,
                "messages": history,
            }
        )

    @router.post("/command/group-chat")
    async def command_group_chat(
        message: str = Form(""),
        conversation_id: str = Form("default"),
        command_user_id: str = Form(""),
        display_name: str = Form(""),
        agent_ids: str = Form(""),
    ) -> JSONResponse:
        cleaned_message = message.strip()
        if not cleaned_message:
            return JSONResponse({"ok": False, "error": "Message required."}, status_code=400)

        cleaned_user_id = _command_user_id(command_user_id)
        group_channel_id = _command_group_channel_id(conversation_id)
        cfg = get_agent_config()
        user_name = display_name.strip() or cfg.get("command_center_name", "Command Center")

        selectable = list_companion_agents("command", selectable_only=True, cfg=cfg)
        selectable_ids = {a["id"] for a in selectable}
        requested = [_normalize_agent_id(a) for a in agent_ids.split(",") if a.strip()]
        group_ids: list[str] = []
        seen: set[str] = set()
        for aid in requested:
            if aid in selectable_ids and aid not in seen:
                seen.add(aid)
                group_ids.append(aid)
        if not group_ids:
            group_ids = [a["id"] for a in selectable]
        if len(group_ids) < 1:
            return JSONResponse({"ok": False, "error": "No companions available for group chat."}, status_code=400)

        agents = [get_companion_agent(aid, cfg) for aid in group_ids]
        replies = await _orchestrate_group_chat(
            agent, cleaned_message, agents, user_name, cleaned_user_id, group_channel_id,
        )
        return JSONResponse(
            {"ok": True, "conversation_id": group_channel_id, "replies": replies}
        )

    @router.get("/command/group-history")
    async def command_group_history(
        conversation_id: str,
        command_user_id: str,
    ) -> JSONResponse:
        cleaned_user_id = _command_user_id(command_user_id)
        group_channel_id = _command_group_channel_id(conversation_id)
        transcript, _ = _load_group(_group_path(cleaned_user_id, group_channel_id))
        messages = [
            {
                "role": "user" if entry.get("type") == "user" else "agent",
                "agent_id": entry.get("agent_id", ""),
                "name": entry.get("name", ""),
                "agent_profile_image": entry.get("profile_image", ""),
                "text": entry.get("text", ""),
            }
            for entry in transcript
        ]
        return JSONResponse(
            {"ok": True, "conversation_id": group_channel_id, "messages": messages}
        )

    @router.post("/command/library-upload")
    async def command_library_upload(
        title: str = Form(""),
        tags: str = Form(""),
        always_on: bool = Form(False),
        platforms: str = Form(""),
        files: list[UploadFile] | None = File(default=None),
    ) -> JSONResponse:
        uploads = files or []
        if not uploads:
            return JSONResponse({"ok": False, "error": "At least one document is required."}, status_code=400)

        created: list[dict[str, Any]] = []
        for index, upload in enumerate(uploads):
            try:
                created.append(
                    await _store_library_module(
                        upload,
                        title=title.strip() if len(uploads) == 1 else "",
                        tags=tags,
                        always_on=always_on,
                        platforms=platforms,
                    )
                )
            except ValueError as exc:
                return JSONResponse(
                    {
                        "ok": False,
                        "error": str(exc),
                        "created": created,
                        "failed_index": index,
                    },
                    status_code=400,
                )

        return JSONResponse({"ok": True, "created": created})

    @router.get("/command/library")
    async def command_library_list() -> JSONResponse:
        store = LibraryStore(str(_ROOT / "data" / "library"))
        return JSONResponse({"ok": True, "modules": store.list_modules()})

    @router.get("/command/library/{module_id}")
    async def command_library_get(module_id: str) -> JSONResponse:
        mod = _get_library_module(module_id)
        if mod is None:
            return JSONResponse({"ok": False, "error": f"Module '{module_id}' not found."}, status_code=404)
        return JSONResponse(
            {
                "ok": True,
                "module": {
                    "id": mod.id,
                    "title": mod.title,
                    "description": mod.description,
                    "always_on": mod.always_on,
                    "platforms": mod.platforms,
                    "tags": mod.tags,
                    "content": mod.content,
                    "char_count": mod.char_count,
                    "filename": mod.filename,
                },
            }
        )

    @router.patch("/command/library/{module_id}")
    async def command_library_patch(module_id: str, body: _LibraryUpdateBody) -> JSONResponse:
        mod = _get_library_module(module_id)
        if mod is None:
            return JSONResponse({"ok": False, "error": f"Module '{module_id}' not found."}, status_code=404)
        _write_library_module(
            filename=mod.filename,
            module_id=mod.id,
            title=mod.title,
            description=mod.description,
            always_on=body.always_on,
            platforms=mod.platforms,
            tags=mod.tags,
            content=mod.content,
        )
        return JSONResponse({"ok": True, "always_on": body.always_on})

    @router.delete("/command/library/{module_id}")
    async def command_library_delete(module_id: str) -> JSONResponse:
        mod = _get_library_module(module_id)
        if mod is None:
            return JSONResponse({"ok": False, "error": f"Module '{module_id}' not found."}, status_code=404)
        path = _ROOT / "data" / "library" / mod.filename
        if path.exists():
            path.unlink()
        return JSONResponse({"ok": True, "deleted": mod.filename})

    return router


async def _build_user_content(message: str, files: list[UploadFile]) -> tuple[str | list[dict[str, Any]], list[dict[str, str]]]:
    blocks: list[dict[str, Any]] = []
    attachments: list[dict[str, str]] = []

    if message:
        blocks.append({"type": "text", "text": message})

    for upload in files:
        raw, media_type, filename = await _read_upload(upload)

        if media_type in _SUPPORTED_IMAGE_MEDIA_TYPES:
            blocks.append({"type": "text", "text": f"Attached image: {filename}"})
            blocks.append(
                {
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": media_type,
                        "data": base64.b64encode(raw).decode("ascii"),
                    },
                }
            )
            attachments.append({"name": filename, "kind": "image", "media_type": media_type})
            continue

        if _is_supported_document(media_type, filename):
            doc_payload = _extract_document_payload(raw, media_type, filename)
            decoded = doc_payload["text"].strip()
            if not decoded:
                raise ValueError(f"{filename} is empty.")
            if len(decoded) > _MAX_TEXT_DOC_CHARS:
                decoded = decoded[:_MAX_TEXT_DOC_CHARS].rstrip() + "\n\n[Document truncated for chat upload.]"
            blocks.append(
                {
                    "type": "text",
                    "text": f"Attached document: {filename}\nMedia type: {media_type}\n\n{decoded}",
                }
            )
            for image in doc_payload.get("images", [])[:_MAX_EMBEDDED_IMAGES]:
                blocks.append({"type": "text", "text": f"Embedded image from {filename}: {image['label']}"})
                blocks.append(
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": image["media_type"],
                            "data": image["data"],
                        },
                    }
                )
            attachments.append(
                {
                    "name": filename,
                    "kind": "document",
                    "media_type": media_type,
                    "images": str(len(doc_payload.get("images", []))),
                    "tables": str(doc_payload.get("tables", 0)),
                }
            )
            continue

        raise ValueError(
            f"Unsupported upload type for {filename}. Supported: images plus text-like documents, PDF, and DOCX."
        )

    if not blocks:
        raise ValueError("Message or attachment required.")
    if len(blocks) == 1 and blocks[0].get("type") == "text":
        return str(blocks[0].get("text", "")), attachments
    return blocks, attachments


async def _read_upload(upload: UploadFile) -> tuple[bytes, str, str]:
    raw = await upload.read()
    media_type = upload.content_type or mimetypes.guess_type(upload.filename or "")[0] or "application/octet-stream"
    filename = upload.filename or "upload"
    if len(raw) > _MAX_UPLOAD_BYTES:
        raise ValueError(f"{filename} exceeds the 5 MB upload limit.")
    return raw, media_type, filename


async def _store_library_module(
    upload: UploadFile,
    *,
    title: str,
    tags: str,
    always_on: bool,
    platforms: str,
) -> dict[str, Any]:
    raw, media_type, filename = await _read_upload(upload)
    if not _is_supported_document(media_type, filename):
        raise ValueError(f"{filename} is not supported for library ingest. Use text-like documents, PDF, or DOCX.")

    doc_payload = _extract_document_payload(raw, media_type, filename)
    extracted = doc_payload["text"].strip()
    if not extracted:
        raise ValueError(f"{filename} did not contain readable text.")
    if len(extracted) > _MAX_LIBRARY_DOC_CHARS:
        extracted = extracted[:_MAX_LIBRARY_DOC_CHARS].rstrip() + "\n\n[Document truncated during library ingest.]"

    base_title = title or _human_title(filename)
    module_base = _slugify(Path(filename).stem)
    module_id = module_base
    filename_safe = f"{module_id}.md"
    parsed_tags = [tag for tag in (_clean_fragment(part, fallback="") for part in tags.split(",")) if tag]
    parsed_platforms = [platform for platform in (_clean_fragment(part.lower(), fallback="") for part in platforms.split(",")) if platform]
    description = f"Imported from {filename}"

    library_dir = _ROOT / "data" / "library"
    library_dir.mkdir(parents=True, exist_ok=True)
    output_path = library_dir / filename_safe
    suffix = 2
    while output_path.exists():
        module_id = f"{module_base}-{suffix}"
        filename_safe = f"{module_id}.md"
        output_path = library_dir / filename_safe
        suffix += 1

    image_lines = [
        f"- {image['label']} ({image['media_type']})"
        for image in doc_payload.get("images", [])
    ]
    header_parts = [f"Source file: {filename}"]
    if doc_payload.get("tables"):
        header_parts.append(f"Tables extracted: {doc_payload['tables']}")
    if image_lines:
        header_parts.append("Embedded images:\n" + "\n".join(image_lines))
    body = "\n\n".join(header_parts + [extracted])

    _write_library_module(
        filename=output_path.name,
        module_id=module_id,
        title=base_title,
        description=description,
        always_on=always_on,
        platforms=parsed_platforms,
        tags=parsed_tags,
        content=f"# {base_title}\n\n{body}\n",
    )
    return {
        "id": module_id,
        "title": base_title,
        "filename": output_path.name,
        "char_count": len(extracted),
        "always_on": always_on,
        "platforms": parsed_platforms,
        "tags": parsed_tags,
        "images": len(doc_payload.get("images", [])),
        "tables": doc_payload.get("tables", 0),
    }


def _is_supported_document(media_type: str, filename: str) -> bool:
    if media_type in _SUPPORTED_DOCUMENT_MEDIA_TYPES:
        return True
    return _is_text_document(media_type, filename)


def _is_text_document(media_type: str, filename: str) -> bool:
    if media_type.startswith("text/") or media_type in _TEXT_MEDIA_TYPES:
        return True
    suffix = (Path(filename).suffix or "").lower()
    return suffix in _TEXT_EXTENSIONS


def _extract_document_payload(raw: bytes, media_type: str, filename: str) -> dict[str, Any]:
    suffix = (Path(filename).suffix or "").lower()
    if media_type == "application/pdf" or suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise ValueError("PDF uploads require the pypdf package to be installed.") from exc

        reader = PdfReader(io.BytesIO(raw))
        page_parts: list[str] = []
        images: list[dict[str, str]] = []
        for index, page in enumerate(reader.pages, start=1):
            try:
                page_text = (page.extract_text(extraction_mode="layout") or "").strip()
            except TypeError:
                page_text = (page.extract_text() or "").strip()
            if not page_text:
                page_text = (page.extract_text() or "").strip()
            if page_text:
                page_parts.append(f"## Page {index}\n\n{page_text}")
            for image_index, image in enumerate(getattr(page, "images", []), start=1):
                media_type_guess = _image_media_type(image.name)
                if not media_type_guess:
                    continue
                image_bytes = getattr(image, "data", b"")
                if not image_bytes:
                    continue
                images.append(
                    {
                        "label": f"Page {index} image {image_index}",
                        "media_type": media_type_guess,
                        "data": base64.b64encode(image_bytes).decode("ascii"),
                    }
                )
        return {"text": "\n\n".join(page_parts), "images": images, "tables": 0}

    if media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or suffix == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise ValueError("DOCX uploads require the python-docx package to be installed.") from exc

        doc = Document(io.BytesIO(raw))
        parts: list[str] = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        table_count = 0
        for table in doc.tables:
            rows: list[list[str]] = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    rows.append(cells)
            if not rows:
                continue
            table_count += 1
            parts.append(_render_markdown_table(rows, table_count))

        images: list[dict[str, str]] = []
        seen_names: set[str] = set()
        for rel in doc.part.rels.values():
            target_ref = getattr(rel, "target_ref", "") or ""
            if "/image" not in target_ref:
                continue
            image_part = getattr(rel, "target_part", None)
            if image_part is None:
                continue
            name = Path(getattr(image_part, "partname", f"image-{len(images)+1}")).name
            if name in seen_names:
                continue
            seen_names.add(name)
            media_type_guess = _image_media_type(name)
            if not media_type_guess:
                continue
            images.append(
                {
                    "label": name,
                    "media_type": media_type_guess,
                    "data": base64.b64encode(image_part.blob).decode("ascii"),
                }
            )
        return {"text": "\n\n".join(parts), "images": images, "tables": table_count}

    if _is_text_document(media_type, filename):
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError:
            text = raw.decode("utf-8", errors="replace")
        return {"text": text, "images": [], "tables": 0}

    raise ValueError(f"Unsupported document type for {filename}.")


def _render_markdown_table(rows: list[list[str]], table_number: int) -> str:
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    divider = ["---"] * width
    body_rows = normalized[1:] or []

    lines = [f"## Table {table_number}", _markdown_row(header), _markdown_row(divider)]
    lines.extend(_markdown_row(row) for row in body_rows)
    return "\n".join(lines)


def _markdown_row(values: list[str]) -> str:
    escaped = [value.replace("|", "\\|") for value in values]
    return "| " + " | ".join(escaped) + " |"


def _image_media_type(name: str) -> str:
    guessed = mimetypes.guess_type(name)[0]
    return guessed if guessed in _SUPPORTED_IMAGE_MEDIA_TYPES else ""


def _write_library_module(
    *,
    filename: str,
    module_id: str,
    title: str,
    description: str,
    always_on: bool,
    platforms: list[str],
    tags: list[str],
    content: str,
) -> None:
    library_dir = _ROOT / "data" / "library"
    library_dir.mkdir(parents=True, exist_ok=True)
    path = library_dir / filename
    payload = (
        "---\n"
        f"id: {module_id}\n"
        f"title: {title}\n"
        f"description: {description}\n"
        f"always_on: {'true' if always_on else 'false'}\n"
        f"platforms: {json.dumps(platforms, ensure_ascii=True)}\n"
        f"tags: {json.dumps(tags, ensure_ascii=True)}\n"
        "---\n\n"
        f"{content.strip()}\n"
    )
    path.write_text(payload, encoding="utf-8")


def _get_library_module(module_id: str):
    store = LibraryStore(str(_ROOT / "data" / "library"))
    return store.get_by_id(module_id)


def _human_title(filename: str) -> str:
    stem = Path(filename).stem.replace("_", " ").replace("-", " ").strip()
    return stem.title() or "Imported Document"


def _slugify(value: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", value.strip().lower()).strip("-")
    return slug[:64] or "library-document"


def _clean_fragment(value: str, *, fallback: str) -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9_-]+", "-", value.strip()).strip("-")
    return cleaned[:64] or fallback


def _command_user_id(value: str) -> str:
    cleaned = _clean_fragment(value, fallback="browser")
    return f"command_user_{cleaned}"


def _command_channel_id(value: str) -> str:
    cleaned = _clean_fragment(value, fallback="default")
    return f"command_chat_{cleaned}"


# ------------------------------------------------------------------ group chat

_GROUPS_DIR = _ROOT / "data" / "memory" / "groups"
_GROUP_MAX_TURNS = 8          # hard cap on agent turns per user message (runaway guard)
_GROUP_MAX_PER_AGENT = 2      # an agent may speak at most twice per user message
_GROUP_TRANSCRIPT_CAP = 200   # entries retained in the shared transcript file


def _command_group_channel_id(value: str) -> str:
    cleaned = _clean_fragment(value, fallback="default")
    return f"command_group_{cleaned}"


def _group_path(user_id: str, channel_id: str) -> Path:
    safe = re.sub(r"[^a-zA-Z0-9_-]+", "_", f"{user_id}__{channel_id}")
    return _GROUPS_DIR / f"{safe}.json"


def _load_group(path: Path) -> tuple[list[dict], dict[str, int]]:
    if not path.exists():
        return [], {}
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return [], {}
    if not isinstance(data, dict):
        return [], {}
    transcript = data.get("transcript")
    cursors = data.get("cursors")
    return (
        transcript if isinstance(transcript, list) else [],
        cursors if isinstance(cursors, dict) else {},
    )


def _save_group(path: Path, transcript: list[dict], cursors: dict[str, int]) -> None:
    if len(transcript) > _GROUP_TRANSCRIPT_CAP:
        drop = len(transcript) - _GROUP_TRANSCRIPT_CAP
        transcript = transcript[drop:]
        cursors = {k: max(0, v - drop) for k, v in cursors.items()}
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps({"transcript": transcript, "cursors": cursors}, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )


def _agent_aliases(agent: dict) -> set[str]:
    """Tokens that count as addressing a companion: its name and each configured
    alias/nickname (like the LSL TRIGGER_NAMES), plus @name forms and bare first names."""
    aliases: set[str] = set()
    names = [agent.get("agent_name", "")] + list(agent.get("aliases") or [])
    for raw in names:
        low = str(raw or "").strip().lower()
        if not low:
            continue
        aliases.add(low)                          # full name/alias (may contain spaces)
        aliases.add("@" + "".join(low.split()))   # @namenospace
        first = low.split()[0]
        aliases.add("@" + first)                  # @first
        if len(first) >= 3:
            aliases.add(first)                    # bare first token
    return aliases


def _agent_mention_ids(text: str, agents: list[dict], exclude_id: str = "") -> list[str]:
    """Return the ids of companions named or @mentioned in text (speaker excluded)."""
    hay = text.lower()
    tokens = set(re.findall(r"@?[a-z0-9]+", hay))
    found: list[str] = []
    for a in agents:
        aid = a["id"]
        if aid == exclude_id:
            continue
        for alias in _agent_aliases(a):
            hit = (
                bool(re.search(r"\b" + re.escape(alias) + r"\b", hay))
                if " " in alias
                else alias in tokens
            )
            if hit:
                found.append(aid)
                break
    return found


async def _orchestrate_group_chat(
    agent: AgentCore,
    message: str,
    agents: list[dict],
    user_name: str,
    user_id: str,
    channel_id: str,
) -> list[dict]:
    """Run one user message through the group: mention-routing + cascade.

    Each responder is fed the labeled transcript delta it has not yet seen and
    answered through the normal per-agent pipeline (so group turns persist into
    each companion's own agent-scoped history and feed its long-term memory).
    """
    path = _group_path(user_id, channel_id)
    transcript, cursors = _load_group(path)

    participants_for_prompt = [{"id": user_id, "name": user_name, "type": "user"}] + [
        {
            "id": a["id"],
            "name": a.get("agent_name", "Agent"),
            "type": "agent",
            "aliases": a.get("aliases") or [],
        }
        for a in agents
    ]
    agent_by_id = {a["id"]: a for a in agents}

    transcript.append(
        {
            "speaker_id": user_id,
            "name": user_name,
            "type": "user",
            "text": message,
            "ts": time.strftime("%Y-%m-%dT%H:%M:%S"),
        }
    )

    # When nobody is named, every companion answers once; otherwise only the named ones.
    mentioned = _agent_mention_ids(message, agents)
    queue: deque[str] = deque(mentioned or [a["id"] for a in agents])
    per_agent: dict[str, int] = {a["id"]: 0 for a in agents}
    replies: list[dict] = []
    turns = 0

    while queue and turns < _GROUP_MAX_TURNS:
        rid = queue.popleft()
        companion = agent_by_id.get(rid)
        if companion is None or per_agent[rid] >= _GROUP_MAX_PER_AGENT:
            continue

        start = cursors.get(rid, 0)
        delta = [e for e in transcript[start:] if e.get("speaker_id") != rid]
        cursors[rid] = len(transcript)
        if not delta:
            continue

        delta_text = "\n".join(f"@{e.get('name', '')}: {e.get('text', '')}" for e in delta)
        ctx = MessageContext(
            platform="command",
            user_id=user_id,
            channel_id=channel_id,
            display_name=user_name,
            agent_id=rid,
            group_participants=participants_for_prompt,
        )
        response = await agent.handle_message(delta_text, ctx, skip_rate_limit=True)
        reply_text = (response.text or "").strip()
        if not reply_text:
            continue

        name = companion.get("agent_name", "Agent")
        image = companion.get("agent_profile_image", "")
        transcript.append(
            {
                "speaker_id": rid,
                "agent_id": rid,
                "name": name,
                "type": "agent",
                "profile_image": image,
                "text": reply_text,
                "ts": time.strftime("%Y-%m-%dT%H:%M:%S"),
            }
        )
        cursors[rid] = len(transcript)
        replies.append(
            {"agent_id": rid, "agent_name": name, "agent_profile_image": image, "text": reply_text}
        )
        per_agent[rid] += 1
        turns += 1

        for mid in _agent_mention_ids(reply_text, agents, exclude_id=rid):
            if per_agent.get(mid, 0) < _GROUP_MAX_PER_AGENT and mid not in queue:
                queue.append(mid)

    _save_group(path, transcript, cursors)
    return replies
